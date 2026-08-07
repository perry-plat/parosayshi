from __future__ import annotations

from pathlib import Path
from typing import Iterable

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Dia_Current_Brand_System_and_Public_Asset_Atlas.docx"
ASSET_DIR = ROOT / "_doc-assets"
SCREENSHOTS = ROOT / "screenshots"

PAGE_W = Inches(8.5)
PAGE_H = Inches(11)
CONTENT_W = Inches(6.5)

BLACK = "020204"
INK = "1A1A1A"
SOFT_INK = "5D5D5D"
PAPER = "F8F8F8"
WHITE = "FFFFFF"
YELLOW = "FFDC5C"
PALE_YELLOW = "FFF4C7"
TABLE_HEADER = "E8EEF5"
LINE = "D9D9D9"
BLUE = "2E74B5"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **edges) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        spec = edges.get(edge)
        if not spec:
            continue
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        for key in ("val", "sz", "space", "color"):
            if key in spec:
                node.set(qn(f"w:{key}"), str(spec[key]))


def set_table_width(table, widths: Iterable[float]) -> None:
    widths = list(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(round(widths[idx] * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def keep_with_next(paragraph, value: bool = True) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    node = p_pr.find(qn("w:keepNext"))
    if node is None:
        node = OxmlElement("w:keepNext")
        p_pr.append(node)
    if not value:
        node.set(qn("w:val"), "0")


def add_hyperlink(paragraph, text: str, url: str, color: str = BLUE, underline: bool = True):
    part = paragraph.part
    rid = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)
    run.append(r_pr)
    txt = OxmlElement("w:t")
    txt.text = text
    run.append(txt)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, fld_end])


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def add_picture(paragraph, path: Path, max_width: float = 6.2, max_height: float = 7.1) -> None:
    with Image.open(path) as img:
        w, h = img.size
    width = max_width
    height = width * h / w
    if height > max_height:
        height = max_height
        width = height * w / h
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width), height=Inches(height))
    alt_by_name = {
        "2026-08-03-current-home-hero.png": "Dia homepage hero with large white Dia wordmark and a screaming human cutout on black.",
        "2026-08-03-current-home-scroll-01.png": "Dia homepage showing a designed Friday Brief report inside a browser frame.",
        "2026-08-03-current-home-scroll-02.png": "Dia homepage use-case navigator with a browser screenshot and workday headline.",
        "2026-08-03-reports-campaign-hero.png": "Dia Reports campaign hero staged in a Brooklyn cafe.",
        "2026-08-03-current-work-hero.png": "Dia for Work page with a calm browser product frame and company-ready message.",
        "2026-08-03-current-students-hero.png": "Dia for Students page with pixel icons and colorful study prompt cards.",
        "2026-08-03-current-skills-hero.png": "Dia Skills Gallery hero showing modular skill examples.",
        "2026-08-03-current-release-notes-hero.png": "Dia Weekly issue 35 masthead, roadmap title, and credited artwork.",
        "2024-12-02-wayback-launch-teaser.png": "Archived December 2024 Dia recruiting page with Come build Dia headline.",
        "2025-05-15-wayback-pre-beta.png": "Archived May 2025 Dia pre-beta page with warm gradient and waitlist message.",
        "2025-06-11-wayback-beta-launch-viewport.png": "Archived June 2025 Dia beta hero with Write with your tabs message.",
        "2025-11-20-wayback-home-viewport.png": "Archived November 2025 Dia homepage with download call to action.",
    }
    alt = alt_by_name.get(path.name, path.stem.replace("-", " "))
    for doc_pr in run._r.xpath(".//wp:docPr"):
        doc_pr.set("descr", alt)
        doc_pr.set("title", alt)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_figure(doc: Document, path: Path, caption: str, max_width: float = 6.2, max_height: float = 7.1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    add_picture(p, path, max_width=max_width, max_height=max_height)
    cap = doc.add_paragraph(caption, style="Figure Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_bullet(doc: Document, text: str, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead and text.startswith(bold_lead):
        p.add_run(bold_lead).bold = True
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25


def add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Number")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25


def add_callout(doc: Document, title: str, body: str, fill: str = PALE_YELLOW, title_color: str = INK) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=150, start=200, bottom=150, end=200)
    set_cell_border(cell, top={"val": "single", "sz": "12", "color": YELLOW}, bottom={"val": "nil"}, start={"val": "nil"}, end={"val": "nil"})
    p = cell.paragraphs[0]
    p.style = doc.styles["Callout Title"]
    r = p.add_run(title)
    r.font.color.rgb = RGBColor.from_string(title_color)
    body_p = cell.add_paragraph(body)
    body_p.paragraph_format.space_after = Pt(0)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_stats(doc: Document, stats: list[tuple[str, str]], columns: int = 3) -> None:
    rows = (len(stats) + columns - 1) // columns
    table = doc.add_table(rows=rows, cols=columns)
    widths = [6.5 / columns] * columns
    set_table_width(table, widths)
    for idx in range(rows * columns):
        cell = table.cell(idx // columns, idx % columns)
        set_cell_shading(cell, PAPER)
        set_cell_border(cell, top={"val": "single", "sz": "5", "color": LINE}, bottom={"val": "single", "sz": "5", "color": LINE}, start={"val": "single", "sz": "5", "color": LINE}, end={"val": "single", "sz": "5", "color": LINE})
        set_cell_margins(cell, top=150, start=150, bottom=150, end=150)
        if idx >= len(stats):
            continue
        value, label = stats[idx]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(value)
        r.bold = True
        r.font.size = Pt(12 if columns >= 5 and len(value) >= 8 else (17 if columns >= 5 else 19))
        r.font.color.rgb = RGBColor.from_string(INK)
        p2 = cell.add_paragraph(label)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        for run in p2.runs:
            run.font.size = Pt(7.2 if columns >= 5 else 8.5)
            run.font.color.rgb = RGBColor.from_string(SOFT_INK)


def add_matrix(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, text in enumerate(headers):
        cell = header.cells[i]
        set_cell_shading(cell, TABLE_HEADER)
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(8.5)
    for row_data in rows:
        row = table.add_row()
        for i, text in enumerate(row_data):
            cell = row.cells[i]
            set_cell_shading(cell, WHITE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            r.font.size = Pt(8.5)
            r.font.color.rgb = RGBColor.from_string(INK)
        for cell in row.cells:
            set_cell_border(cell, bottom={"val": "single", "sz": "4", "color": LINE})
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_source(doc: Document, key: str, label: str, url: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"{key}  ")
    r.bold = True
    add_hyperlink(p, label, url)


def new_page(doc: Document, kicker: str, title: str, deck: str | None = None) -> None:
    doc.add_page_break()
    p = doc.add_paragraph(kicker.upper(), style="Kicker")
    keep_with_next(p)
    h = doc.add_heading(title, level=1)
    keep_with_next(h)
    if deck:
        d = doc.add_paragraph(deck, style="Deck")
        keep_with_next(d, False)


def build() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    OUT.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.page_width = PAGE_W
    section.page_height = PAGE_H
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    props = doc.core_properties
    props.title = "Dia — Current Brand System & Public Asset Atlas"
    props.subject = "Public-web design research, current-state weighted"
    props.author = "Paro / Codex research collaboration"
    props.comments = "Base preset: compact_reference_guide. Named Dia accent override: near-black headings, warm-white grounds, yellow callouts, editorial cover."
    props.keywords = "Dia, Browser Company, brand, design system, asset archive, typography, voice, release notes"

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor.from_string(INK)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor.from_string(INK)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor.from_string(INK)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(5)
    h3.paragraph_format.keep_with_next = True

    for base_name in ("List Bullet", "List Number"):
        st = doc.styles[base_name]
        st.font.name = "Calibri"
        st.font.size = Pt(11)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.25

    cover_title = doc.styles.add_style("Cover Title", WD_STYLE_TYPE.PARAGRAPH)
    cover_title.font.name = "Georgia"
    cover_title.font.size = Pt(31)
    cover_title.font.bold = True
    cover_title.font.color.rgb = RGBColor.from_string(WHITE)
    cover_title.paragraph_format.space_after = Pt(12)

    cover_deck = doc.styles.add_style("Cover Deck", WD_STYLE_TYPE.PARAGRAPH)
    cover_deck.font.name = "Calibri"
    cover_deck.font.size = Pt(12)
    cover_deck.font.color.rgb = RGBColor.from_string("D9D9D9")
    cover_deck.paragraph_format.space_after = Pt(12)
    cover_deck.paragraph_format.line_spacing = 1.18

    kicker = doc.styles.add_style("Kicker", WD_STYLE_TYPE.PARAGRAPH)
    kicker.font.name = "Courier New"
    kicker.font.size = Pt(8)
    kicker.font.bold = True
    kicker.font.color.rgb = RGBColor.from_string(SOFT_INK)
    kicker.paragraph_format.space_after = Pt(5)

    deck = doc.styles.add_style("Deck", WD_STYLE_TYPE.PARAGRAPH)
    deck.font.name = "Georgia"
    deck.font.size = Pt(12)
    deck.font.italic = True
    deck.font.color.rgb = RGBColor.from_string(SOFT_INK)
    deck.paragraph_format.space_after = Pt(12)
    deck.paragraph_format.line_spacing = 1.15

    fig = doc.styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    fig.font.name = "Calibri"
    fig.font.size = Pt(8.5)
    fig.font.italic = True
    fig.font.color.rgb = RGBColor.from_string(SOFT_INK)
    fig.paragraph_format.space_after = Pt(8)
    fig.paragraph_format.keep_with_next = False

    callout_title = doc.styles.add_style("Callout Title", WD_STYLE_TYPE.PARAGRAPH)
    callout_title.font.name = "Courier New"
    callout_title.font.size = Pt(8.5)
    callout_title.font.bold = True
    callout_title.font.color.rgb = RGBColor.from_string(INK)
    callout_title.paragraph_format.space_after = Pt(4)

    # Interior header and footer; the cover remains clean.
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("DIA / PUBLIC BRAND REFERENCE / 2026-08-03")
    hr.font.name = "Courier New"
    hr.font.size = Pt(7.5)
    hr.font.color.rgb = RGBColor.from_string("777777")

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = fp.add_run("PUBLIC-WEB RESEARCH  ·  ")
    fr.font.name = "Courier New"
    fr.font.size = Pt(7.5)
    fr.font.color.rgb = RGBColor.from_string("777777")
    add_page_number(fp)

    # Cover — editorial_cover template with a Dia-accent color treatment.
    cover = doc.add_table(rows=1, cols=2)
    set_table_width(cover, [3.05, 3.45])
    cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    row = cover.rows[0]
    row.height = Inches(7.75)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    left, right = row.cells
    for cell in row.cells:
        set_cell_shading(cell, BLACK)
        set_cell_border(cell, top={"val": "nil"}, bottom={"val": "nil"}, start={"val": "nil"}, end={"val": "nil"})
        set_cell_margins(cell, top=260, start=260, bottom=260, end=260)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = left.paragraphs[0]
    p.style = doc.styles["Kicker"]
    r = p.add_run("CURRENT-STATE WEIGHTED / PUBLIC WEB")
    r.font.color.rgb = RGBColor.from_string(YELLOW)
    t = left.add_paragraph(style="Cover Title")
    t.add_run("Dia\nCurrent Brand System\n& Public Asset Atlas")
    d = left.add_paragraph(style="Cover Deck")
    d.add_run("A working reference for visual language, voice, assets, motion, product storytelling, and historical lineage.")
    meta = left.add_paragraph()
    mr = meta.add_run("OBSERVED 03 AUG 2026\nCURRENT 78% / LINEAGE 22%")
    mr.font.name = "Courier New"
    mr.font.size = Pt(8)
    mr.font.color.rgb = RGBColor.from_string("BDBDBD")
    rp = right.paragraphs[0]
    add_picture(rp, SCREENSHOTS / "2026-08-03-current-home-hero.png", max_width=3.0, max_height=6.7)
    accent = doc.add_table(rows=1, cols=1)
    set_table_width(accent, [6.5])
    set_cell_shading(accent.cell(0, 0), YELLOW)
    accent.rows[0].height = Inches(0.08)
    accent.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    # Page 2 — scope and method.
    new_page(doc, "How to read this", "A current brand guide with a forensic appendix", "The main body explains Dia as it exists now. The historical section is retained to show which ideas persisted, which were discarded, and where the present system came from.")
    add_callout(doc, "THE SHORT READ", "Dia has moved from explaining an AI browser to dramatizing a better workday. The brand now makes a conspicuous emotional promise, then proves it through calm editorial outputs, browser-native context, and an unusually authored release culture.")
    doc.add_heading("Collection boundary", level=2)
    add_bullet(doc, "Current public HTML and CSS across the main brand, product, work, students, skills, security, release-notes, Windows, and Reports surfaces.")
    add_bullet(doc, "Public Sanity content and asset records, including images, Mux video references, downloadable animation files, skills, packs, and release systems.")
    add_bullet(doc, "A deduplicated Wayback capture index from 2024–2026 plus selected rendered moments at major brand transitions.")
    add_bullet(doc, "A filtered language corpus of brand pages and release communication, excluding the internal-looking public test article.")
    doc.add_heading("What this is not", level=2)
    p = doc.add_paragraph("This is not an official brand book and it does not grant reuse rights. It cannot recover authenticated, unpublished, privately held, deleted, or robots-blocked material. Fonts and third-party art are identified for reference; licensed binaries are not redistributed.")
    p.paragraph_format.space_after = Pt(10)
    add_stats(doc, [("304", "public content documents"), ("875", "public image-asset records"), ("263", "Mux video-asset records"), ("7,902", "deduplicated Wayback records"), ("563", "downloaded image previews"), ("~395 MB", "reference archive on disk")])

    # Page 3 — current position.
    new_page(doc, "01 / Now", "The current brand in one sentence", "Dia sells relief from the fragmented workday: the browser carries context across tabs, turns it into useful synthesis, and gives the user a calmer sense of readiness.")
    doc.add_heading("The present hierarchy", level=2)
    rows = [
        ["1", "Emotional tension", "Dread, hunting, prep, fragmentation, doing it alone."],
        ["2", "Browser-native advantage", "Dia can see the context already distributed across tabs and work surfaces."],
        ["3", "Useful synthesis", "Briefs, reports, notes, answers, memory, and meeting readiness."],
        ["4", "Human resolution", "Start ahead; show up; move faster; have a better day."],
        ["5", "Trust layer", "Privacy, control, security, company readiness, and clear provenance."],
    ]
    add_matrix(doc, ["ORDER", "JOB", "CURRENT EXPRESSION"], rows, [0.65, 1.75, 4.1])
    doc.add_heading("Why it feels different from generic AI marketing", level=2)
    add_bullet(doc, "The homepage leads with emotion rather than model capability.", "The homepage")
    add_bullet(doc, "AI is explained through continuity and outputs, not spectacle or futurist vocabulary.", "AI")
    add_bullet(doc, "The interface is quiet enough to make the authored content—not the chrome—the hero.", "The interface")
    add_bullet(doc, "Dia repeatedly describes ordinary work: tabs, meetings, groups, reports, context, and the day ahead.", "Dia")
    add_callout(doc, "CURRENT BRAND AXIS", "Raw emotion ↔ calm competence. Dia earns memorability with the first half and trust with the second.")

    # Page 4 — opening system.
    new_page(doc, "02 / Homepage", "One loud idea, then authored proof", "The current homepage behaves like an emotional cold open followed by an editorial product demonstration.")
    add_figure(doc, SCREENSHOTS / "2026-08-03-current-home-scroll-01.png", "Figure 1. Current homepage after the hero: an authored “Friday Brief” inside browser chrome. Source: current Dia homepage, observed 2026-08-03.", max_width=4.75, max_height=5.95)
    p = doc.add_paragraph()
    p.add_run("Read: ").bold = True
    p.add_run("The product claim becomes credible when it is embodied as a designed artifact. The landscape illustration, vertical date, report typography, and browser frame make synthesis feel authored rather than autogenerated.")

    # Page 5 — scroll choreography.
    new_page(doc, "02 / Homepage", "The scroll choreography is the argument", "Dia stages the product as a sequence of situations, not a grid of equal-weight features.")
    add_figure(doc, SCREENSHOTS / "2026-08-03-current-home-scroll-02.png", "Figure 2. Current homepage use-case navigator: one active story, one product frame, one transition toward the broader work promise.", max_width=5.2, max_height=4.6)
    doc.add_heading("Observed sequence", level=2)
    add_numbered(doc, "A black, emotionally extreme hero makes the brand impossible to ignore.")
    add_numbered(doc, "A pale editorial section reframes Dia as a thoughtful interpreter of browser context.")
    add_numbered(doc, "A controlled use-case sequence changes one example at a time inside a stable browser frame.")
    add_numbered(doc, "The page zooms back out to the universal promise: a browser built around the user’s actual work.")
    add_callout(doc, "TRANSFERABLE PRINCIPLE", "Spend one special effect per viewport. Keep the product evidence compositionally quiet enough to be read.")

    # Page 6 — type.
    new_page(doc, "03 / Visual system", "Typography separates feeling, function, and metadata", "The current system uses type roles more strategically than a conventional display/body pairing.")
    type_rows = [
        ["Exposure Variable / Exposure VAR", "Editorial display", "Hero wordmark, emotional headlines, Reports, release-note titles", "Highly expressive; unusual exposure axis and many static instances"],
        ["ABC Oracle", "Product sans", "Body, work pages, product explanation, many current headings", "Quiet, humanist, low-friction"],
        ["ABC Favorit Mono", "Utility mono", "Navigation, footer, compact labels", "Operational and browser-like"],
        ["Quadrant Text Mono", "Editorial metadata", "Release issue, version, date, location", "Makes shipping feel archived and collectible"],
        ["ABC Oracle Triple", "Accent display", "Selected campaign or component accents", "Use sparingly"],
        ["SF Pro / system fallbacks", "Interface support", "Product-adjacent UI and resilience", "Familiar platform texture"],
    ]
    add_matrix(doc, ["FAMILY", "ROLE", "OBSERVED USE", "EFFECT"], type_rows, [1.55, 1.15, 2.15, 1.65])
    doc.add_heading("Current scale cues", level=2)
    add_bullet(doc, "Home hero “Dia”: approximately 112 px on desktop; the word behaves as image and name at once.")
    add_bullet(doc, "Major editorial sections: approximately 48 px Exposure; major product sections: approximately 54 px ABC Oracle.")
    add_bullet(doc, "Product descriptions: approximately 18 px ABC Oracle with softened black.")
    add_bullet(doc, "Release-note metadata and footer systems: approximately 13 px mono, often spaced and boxed.")
    add_callout(doc, "LICENSING", "Do not copy or redistribute these commercial fonts from the reference archive. Recreate the role contrast with licensed or open alternatives.", fill="FCE8E6")

    # Page 7 — color/material/motion.
    new_page(doc, "03 / Visual system", "The base is nearly colorless; the artifacts carry the chroma", "Dia’s core page system is near-black, warm white, soft gray, and glassy browser chrome. Strong color arrives inside campaigns, reports, skills, and illustrations.")
    palette = doc.add_table(rows=2, cols=5)
    set_table_width(palette, [1.3] * 5)
    colors = [("#020204", BLACK), ("#F8F8F8", PAPER), ("#FFFFFF", WHITE), ("#D9D9D9", LINE), ("#FFDC5C", YELLOW)]
    for idx, (label, fill) in enumerate(colors):
        cell = palette.cell(0, idx)
        set_cell_shading(cell, fill)
        cell.height = Inches(0.65)
        set_cell_border(cell, top={"val": "single", "sz": "4", "color": LINE}, bottom={"val": "single", "sz": "4", "color": LINE}, start={"val": "single", "sz": "4", "color": LINE}, end={"val": "single", "sz": "4", "color": LINE})
        lab = palette.cell(1, idx)
        p = lab.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = p.add_run(label)
        rr.font.name = "Courier New"
        rr.font.size = Pt(8)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    doc.add_heading("Material grammar", level=2)
    add_bullet(doc, "Pill-shaped navigation and controls echo browser chrome without reproducing a literal toolbar.")
    add_bullet(doc, "Large radii, pale glow fields, thin gray rules, and sparse shadows keep surfaces soft rather than glossy.")
    add_bullet(doc, "Browser screenshots are framed as finished scenes; the surrounding page rarely competes with the screen.")
    add_bullet(doc, "Paper, postcards, painting, mono stamps, clips, and vertical metadata turn digital updates into editorial objects.")
    doc.add_heading("Motion grammar", level=2)
    add_bullet(doc, "Pinned or staged scroll transitions hold one composition while the active narrative changes.")
    add_bullet(doc, "The CSS includes authored easing families such as cubic-bezier(.76,0,.24,1) and cubic-bezier(.87,0,.13,1), alongside standard product easings.")
    add_bullet(doc, "Motion reveals state and attention; it is not used as permanent ambient decoration.")

    # Page 8 — voice corpus.
    new_page(doc, "04 / Language", "The voice is direct, second-person, and emotionally ordinary", "Dia now talks like a thoughtful colleague who already knows the situation—not a lab explaining a breakthrough.")
    add_stats(doc, [("26,412", "words in filtered corpus"), ("1,271", "second-person uses"), ("294", "first-person-plural uses"), ("656", "contractions"), ("92", "questions"), ("127", "em dashes")])
    doc.add_heading("What the counts suggest", level=2)
    add_bullet(doc, "Second-person language appears about 4.3× as often as first-person plural: the user’s situation is the grammatical center.")
    add_bullet(doc, "Contractions keep the writing spoken and contemporary without becoming internet-slang heavy.")
    add_bullet(doc, "Repeated concepts include tabs, work, chat, groups, meetings, context, open, day, without, and better.")
    add_bullet(doc, "Questions are used to activate recognition and frame a job, not to manufacture mystery.")
    doc.add_heading("Current copy ladder", level=2)
    copy_rows = [
        ["Emotional hook", "A browser you won’t dread opening.", "Names the feeling before the feature."],
        ["Mechanism", "Dia reads between the tabs.", "Makes context continuity legible."],
        ["Fit", "Built for how you actually work.", "Rejects generic productivity theater."],
        ["Resolution", "Ready for a better day?", "Returns the product promise to lived experience."],
    ]
    add_matrix(doc, ["LAYER", "PUBLIC EXAMPLE", "JOB"], copy_rows, [1.25, 2.45, 2.8])
    add_callout(doc, "COPY TEST", "If a line could describe any AI tool, it is not yet Dia-shaped. Anchor the line in a browser situation, an authored output, or a familiar workday feeling.")

    # Page 9 — voice rules.
    new_page(doc, "04 / Language", "A usable writing formula", "The current voice can be replicated as a decision sequence without copying any sentence.")
    formula = [
        ("1", "Name the friction", "A human feeling or ordinary task: hunting, preparing, catching up, doing it alone."),
        ("2", "Reveal the contextual advantage", "The browser already contains the tabs, history, people, and work needed to help."),
        ("3", "Materialize the intelligence", "Show the brief, report, summary, answer, or next step."),
        ("4", "Resolve in life", "Ahead, ready, focused, calmer, moving, or done."),
        ("5", "Prove control", "When stakes rise, explain privacy, choice, source scope, and company readiness."),
        ("6", "Invite lightly", "Download, try, ask, explore, or add it to the day."),
    ]
    add_matrix(doc, ["STEP", "MOVE", "WHAT TO WRITE"], [list(x) for x in formula], [0.6, 1.55, 4.35])
    doc.add_heading("Do", level=2)
    add_bullet(doc, "Use familiar nouns and verbs before technical terms.")
    add_bullet(doc, "Let one sentence carry one job.")
    add_bullet(doc, "Write toward a concrete output or a changed state.")
    add_bullet(doc, "Use contractions and direct address when they improve rhythm.")
    doc.add_heading("Avoid", level=2)
    add_bullet(doc, "Generic superlatives, model worship, or claims that require trust before evidence.")
    add_bullet(doc, "Feature pile-ups, autonomous-agent theater, or vague “transform your workflow” language.")
    add_bullet(doc, "Permanent whimsy. Dia’s playful or raw gestures sit inside a disciplined system.")

    # Page 10 — current surfaces.
    new_page(doc, "05 / Surface map", "One brand, several levels of intensity", "The current ecosystem does not use a single tone everywhere. Each surface carries a specific proof burden.")
    surface_rows = [
        ["Homepage", "Emotional brand platform", "Raw opening; calm authored proof", "Dread → context → better day"],
        ["Reports / Start", "Campaign and routine", "Documentary, local, warm", "Bring synthesis into the morning"],
        ["Dia Weekly", "Release culture", "Editorial, credited, collectible", "Shipping is a publication"],
        ["Work", "Company readiness", "Quiet, broad, reassuring", "Context plus security and scale"],
        ["Students", "High-energy utility", "Playful, direct, graphic", "You do not have to do it alone"],
        ["Skills", "Extensibility", "Gallery-like, modular", "See possibilities and reusable packs"],
        ["Windows", "Platform expansion", "Restrained beta utility", "Join and stay informed"],
        ["Security / Privacy", "Trust infrastructure", "Specific and controlled", "User agency and evidence"],
    ]
    add_matrix(doc, ["SURFACE", "ROLE", "ENERGY", "PROMISE"], surface_rows, [1.2, 1.55, 1.55, 2.2])
    add_callout(doc, "SYSTEM RULE", "Intensity belongs to the surface’s job. The homepage can scream; the company page should reassure; the release system can become an authored magazine.")

    # Page 11 — reports and work proof.
    new_page(doc, "05 / Surface map", "Campaign worlds make the product socially believable", "Reports is staged in a real café and a real morning. Work is staged as calm product evidence. Together they connect emotional brand to practical adoption.")
    table = doc.add_table(rows=1, cols=2)
    set_table_width(table, [3.25, 3.25])
    for idx, (path, title) in enumerate([
        (SCREENSHOTS / "2026-08-03-reports-campaign-hero.png", "REPORTS / DOCUMENTARY ROUTINE"),
        (SCREENSHOTS / "2026-08-03-current-work-hero.png", "WORK / QUIET PROOF"),
    ]):
        cell = table.cell(0, idx)
        set_cell_border(cell, top={"val": "nil"}, bottom={"val": "nil"}, start={"val": "nil"}, end={"val": "nil"})
        set_cell_margins(cell, top=60, start=60, bottom=60, end=60)
        p = cell.paragraphs[0]
        add_picture(p, path, max_width=3.0, max_height=4.0)
        c = cell.add_paragraph(title)
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in c.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(7.5)
            run.font.color.rgb = RGBColor.from_string(SOFT_INK)
    doc.add_heading("Design implication", level=2)
    p = doc.add_paragraph("Dia does not ask one visual system to carry every level of proof. It alternates between authored cultural settings and precise product frames, then lets a consistent type hierarchy and navigation system hold the family together.")

    # Page 12 — students and skills.
    new_page(doc, "05 / Surface map", "Audience pages are allowed to change costume", "Students uses pixel icons and high-chroma prompt cards. Skills uses a gallery logic. Both remain recognizable through type roles, spacing, chrome, and direct utility copy.")
    table2 = doc.add_table(rows=1, cols=2)
    set_table_width(table2, [3.25, 3.25])
    for idx, (path, title) in enumerate([
        (SCREENSHOTS / "2026-08-03-current-students-hero.png", "STUDENTS / PLAYFUL ASSISTANCE"),
        (SCREENSHOTS / "2026-08-03-current-skills-hero.png", "SKILLS / MODULAR POSSIBILITY"),
    ]):
        cell = table2.cell(0, idx)
        set_cell_border(cell, top={"val": "nil"}, bottom={"val": "nil"}, start={"val": "nil"}, end={"val": "nil"})
        set_cell_margins(cell, top=60, start=60, bottom=60, end=60)
        p = cell.paragraphs[0]
        add_picture(p, path, max_width=3.0, max_height=4.0)
        c = cell.add_paragraph(title)
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in c.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(7.5)
            run.font.color.rgb = RGBColor.from_string(SOFT_INK)
    add_callout(doc, "BRAND RANGE", "Consistency is structural, not cosmetic: type roles, direct address, browser-native context, and authored outputs survive even when the illustration style changes.")

    # Page 13 — release editorial system.
    new_page(doc, "06 / Release culture", "Dia Weekly turns shipping into editorial memory", "The release system is the strongest example of Dia’s brand language becoming an operating system rather than a campaign.")
    add_figure(doc, SCREENSHOTS / "2026-08-03-current-release-notes-hero.png", "Figure 3. Dia Weekly Unfiltered, issue 035 / app 1.42.0: publication masthead, date, time, location, version, art, and a story-led title.", max_width=4.92, max_height=4.35)
    doc.add_heading("Recurring anatomy", level=2)
    add_bullet(doc, "Issue number, app version, release date/time, and location establish provenance.")
    add_bullet(doc, "Story-led titles replace purely functional version labels.")
    add_bullet(doc, "Builder credits, teammate modules, postcards, art, and personal cultural recommendations make the company visible.")
    add_bullet(doc, "Wide media, showcases, lists, grids, and notes let each issue have a composed rhythm.")
    add_stats(doc, [("35", "issues"), ("2025-11-11", "first issue"), ("2026-07-30", "latest"), ("80", "video refs"), ("215", "image uses"), ("22", "legacy docs")], columns=6)

    # Page 14 — image language.
    new_page(doc, "07 / Image language", "Five image families carry most of the brand", "The archive is large, but the current imagery becomes manageable when grouped by the job each image performs.")
    image_rows = [
        ["Emotional human", "Scream cutout and expressive figures", "Instant memory; names felt friction"],
        ["Documentary setting", "Café, team, morning, city, real rooms", "Places AI inside lived routine"],
        ["Authored output", "Reports, briefs, notes, study cards", "Makes synthesis visible and desirable"],
        ["Product proof", "Browser frames, dialogs, groups, skills", "Shows the mechanism without overexplaining"],
        ["Editorial ephemera", "Postcards, paper, clips, painting, mono stamps", "Gives release history provenance and culture"],
    ]
    add_matrix(doc, ["FAMILY", "EXAMPLES", "BRAND JOB"], image_rows, [1.35, 2.35, 2.8])
    doc.add_heading("Selection rules", level=2)
    add_bullet(doc, "Prefer images that establish a situation or output, not generic visual mood.")
    add_bullet(doc, "Let the dominant image have a clear editorial crop and enough quiet space around it.")
    add_bullet(doc, "Credit artists and builders when the release system treats the work as authored.")
    add_bullet(doc, "Keep high-chroma imagery inside a neutral frame so color reads as content, not decoration.")
    add_bullet(doc, "Avoid mixing multiple novelty styles in the same viewport.")
    add_callout(doc, "ARCHIVE PRACTICE", "Use the 563 preview images to browse visual families. Use images.csv to recover dimensions, IDs, original URLs, and usage context before choosing a production reference.")

    # Page 15 — Parosayshi translation.
    new_page(doc, "08 / For ours", "Borrow the grammar, not the skin", "Dia is useful to Parosayshi because both can turn product history into authored physical evidence. The opportunity is translation, not imitation.")
    translate_rows = [
        ["Expressive cold open", "One tactile or human gesture that names the project’s tension", "Do not copy the scream or black hero literally"],
        ["Editorial output", "A proof card, field note, receipt, report, or annotated artifact", "Use real case evidence, not ornamental filler"],
        ["Mono provenance", "Issue/date/version/role stamps on case-study objects", "Keep it subordinate to the tactile object"],
        ["Calm browser proof", "Stable reader or device frame for decision evidence", "Avoid over-polished SaaS mockup sameness"],
        ["Release culture", "Case notes and process drops that accumulate over time", "Credit ownership and distinguish fact from contribution"],
        ["One effect per viewport", "Let a special interaction own the moment, then settle", "Preserve the cutting mat, paper, folders, and reader"],
    ]
    add_matrix(doc, ["DIA GRAMMAR", "PAROSAYSHI TRANSLATION", "GUARDRAIL"], translate_rows, [1.45, 2.8, 2.25])
    doc.add_heading("The strongest existing bridge", level=2)
    p = doc.add_paragraph("The earlier Dia-inspired postcard direction already translated release-note metadata—issue, date, version, authorship—into a physical portfolio object. The next useful step is to turn that from a one-off reference into a reusable provenance layer for selected case studies, while keeping the existing material world dominant.")
    add_callout(doc, "DECISION", "Use Dia to sharpen hierarchy, provenance, and authored proof. Do not use it to replace Parosayshi’s tactile identity with a generic AI-browser aesthetic.")

    # Page 16 — condensed history.
    new_page(doc, "09 / Lineage", "The brand moved from possibility to utility to emotional relief", "History matters because the current system is not a visual reskin; it changes what Dia believes needs to be sold.")
    hist_rows = [
        ["2024-12", "Recruiting / vision", "Warm beige, GT Ultra Median, minimal page", "Come build Dia"],
        ["2025-05", "Pre-beta / waitlist", "Warm restraint; early audience activation", "Stay in the loop"],
        ["2025-06", "Beta explanation", "ABC Oracle, job-based sections, testimonials", "Chat with your tabs"],
        ["2025-11", "General availability", "Broader jobs-to-be-done; Arc provenance", "Dia. Do it all."],
        ["2025-11 → 2026-07", "Editorial shipping", "Dia Weekly, issue/version metadata, credits", "Release notes as publication"],
        ["2026-08", "Workday relief", "Scream hero, reports, better-day narrative", "Emotion → synthesis → readiness"],
    ]
    add_matrix(doc, ["MOMENT", "BRAND TASK", "VISUAL / SYSTEM CUE", "MESSAGE"], hist_rows, [1.05, 1.55, 2.45, 1.45])
    doc.add_heading("What persisted", level=2)
    add_bullet(doc, "Editorial serif authority, direct writing, and a belief that the browser can become a personal context layer.")
    add_bullet(doc, "Warmth and humanity, even as the palette moved toward black, white, and more authored campaign worlds.")
    add_bullet(doc, "A consistent refusal to make technical AI language the emotional center of the brand.")
    doc.add_heading("What changed", level=2)
    add_bullet(doc, "The promise narrowed from doing many things to feeling better prepared for the actual day.")
    add_bullet(doc, "Feature explanation matured into output-led storytelling and an editorial release institution.")

    # Page 17 — historical visuals.
    new_page(doc, "09 / Lineage", "Selected archived states", "These are representative pivots, not every Wayback frame. The complete deduplicated capture index is retained in the archive.")
    table3 = doc.add_table(rows=2, cols=2)
    set_table_width(table3, [3.25, 3.25])
    historical = [
        (SCREENSHOTS / "2024-12-02-wayback-launch-teaser.png", "2024-12 / recruiting reveal"),
        (SCREENSHOTS / "2025-05-15-wayback-pre-beta.png", "2025-05 / pre-beta"),
        (SCREENSHOTS / "2025-06-11-wayback-beta-launch-viewport.png", "2025-06 / beta launch"),
        (SCREENSHOTS / "2025-11-20-wayback-home-viewport.png", "2025-11 / general availability"),
    ]
    for idx, (path, label) in enumerate(historical):
        cell = table3.cell(idx // 2, idx % 2)
        set_cell_border(cell, top={"val": "nil"}, bottom={"val": "nil"}, start={"val": "nil"}, end={"val": "nil"})
        set_cell_margins(cell, top=60, start=60, bottom=80, end=60)
        p = cell.paragraphs[0]
        add_picture(p, path, max_width=3.0, max_height=2.3)
        c = cell.add_paragraph(label.upper())
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in c.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(7.5)
            run.font.color.rgb = RGBColor.from_string(SOFT_INK)
    add_callout(doc, "HISTORY WEIGHT", "The guide uses these states to explain lineage. It does not recommend returning to their warm-beige minimalism or their feature-by-feature launch structure.")

    # Page 18 — archive map.
    new_page(doc, "10 / Archive", "What was captured and how to use it", "The archive is designed for browsing, tracing, and refreshing—not just storing screenshots.")
    archive_rows = [
        ["previews/images/", "563 current content-linked previews", "Visual browsing and contact-sheet work"],
        ["manifests/images.csv", "875 image asset records", "IDs, dimensions, sizes, URLs, filenames"],
        ["manifests/videos.csv", "263 Mux video records / 2h 15m 26s", "Stream and poster reference; motion audit"],
        ["files/ + manifests/files.csv", "16 public animation JSON files", "Inspect motion assets and implementation"],
        ["raw/sanity/", "304 content docs plus asset records", "Trace images and copy back to content"],
        ["raw/site/", "Current public HTML and CSS", "Typography, color, layout, route, and code signals"],
        ["raw/wayback/", "7,902 deduplicated capture records", "Historical discovery and missing-state lookup"],
        ["analysis/", "Copy corpus, records, stats, current brief", "Language analysis and design synthesis"],
        ["tools/collect-dia.mjs", "Repeatable public collector", "Refresh the current snapshot later"],
    ]
    add_matrix(doc, ["PATH", "CONTENTS", "BEST USE"], archive_rows, [1.75, 2.35, 2.4])
    doc.add_heading("Rights-aware handling", level=2)
    add_bullet(doc, "Treat all Dia, artist, photographer, and product assets as references unless permission is confirmed.")
    add_bullet(doc, "Use the font endpoint catalog to identify families; do not copy the commercial binaries.")
    add_bullet(doc, "Prefer transformed previews for research and recover original URLs from the manifests only when needed.")
    add_bullet(doc, "Re-run the collector before a consequential design decision; the public corpus is live and can change.")

    # Page 19 — sources.
    new_page(doc, "11 / Sources", "Primary source atlas", "Official current pages, public structured data, and timestamped Wayback states form the evidence base. Company context is limited to official announcements.")
    doc.add_heading("Current official surfaces", level=2)
    add_source(doc, "S1", "Dia homepage", "https://www.diabrowser.com/")
    add_source(doc, "S2", "Dia Weekly / latest release notes", "https://www.diabrowser.com/release-notes/latest")
    add_source(doc, "S3", "Introducing Reports / Start", "https://www.diabrowser.com/start")
    add_source(doc, "S4", "Dia for Work", "https://www.diabrowser.com/forwork")
    add_source(doc, "S5", "Dia for Students", "https://www.diabrowser.com/students")
    add_source(doc, "S6", "Dia Skills Gallery", "https://www.diabrowser.com/skills")
    add_source(doc, "S7", "Dia Security", "https://www.diabrowser.com/security")
    add_source(doc, "S8", "Dia Privacy", "https://www.diabrowser.com/privacy")
    doc.add_heading("Official history and company context", level=2)
    add_source(doc, "S9", "Early peek at Dia — recruiting video", "https://www.youtube.com/watch?v=C25g53PC5QQ")
    add_source(doc, "S10", "Letter to Arc members — The Browser Company", "https://browsercompany.substack.com/p/letter-to-arc-members-2025")
    add_source(doc, "S11", "Atlassian acquisition announcement", "https://www.atlassian.com/blog/announcements/atlassian-acquires-the-browser-company")
    doc.add_heading("Selected Wayback states", level=2)
    add_source(doc, "W1", "2024-12-02 Dia public reveal", "https://web.archive.org/web/20241202145719id_/https://www.diabrowser.com/")
    add_source(doc, "W2", "2025-05-15 pre-beta", "https://web.archive.org/web/20250515024755id_/https://www.diabrowser.com/")
    add_source(doc, "W3", "2025-06-11 beta launch", "https://web.archive.org/web/20250611184246id_/https://www.diabrowser.com/")
    add_source(doc, "W4", "2025-11-20 general-availability-era home", "https://web.archive.org/web/20251120012357id_/https://www.diabrowser.com/")
    add_callout(doc, "OBSERVATION DATE", "Current-page claims, counts, and screenshots in this guide were observed or collected on 2026-08-03. Refresh before treating them as present-day facts.")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
